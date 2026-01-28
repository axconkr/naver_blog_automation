"""
워드 문서 생성 모듈

엑셀에서 블로그 본문을 읽어 워드 문서로 변환하고,
각 섹션에 AI 생성 이미지를 추가합니다.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
from nanobanana import generate_image, generate_blog_images
from dotenv import load_dotenv
import os
import glob
import re
import time

load_dotenv()

# 출력 디렉토리
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
IMAGES_DIR = os.path.join(os.path.dirname(__file__), "images")


def find_blog_file():
    """현재 디렉토리에서 blog+날짜.xlsx 파일 찾기"""
    current_dir = os.getcwd()
    pattern = os.path.join(current_dir, "blog*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        raise FileNotFoundError("blog+날짜.xlsx 파일을 찾을 수 없습니다.")
    
    return max(files, key=os.path.getmtime)


def parse_blog_content(content):
    """
    블로그 본문을 섹션별로 파싱
    
    Returns:
        dict: {
            "title": "제목",
            "summary": "요약",
            "sections": [{"heading": "소제목", "content": "본문"}, ...],
            "faq": [{"question": "Q", "answer": "A"}, ...],
            "hashtags": ["#태그1", "#태그2", ...]
        }
    """
    result = {
        "title": "",
        "summary": "",
        "sections": [],
        "faq": [],
        "hashtags": []
    }
    
    lines = content.strip().split('\n')
    
    current_section = None
    current_content = []
    in_faq = False
    current_question = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 제목 파싱 (첫 번째 줄 또는 [제목:...] 형식)
        if line.startswith('[제목:') or line.startswith('제목:'):
            title = re.sub(r'^\[?제목:\s*', '', line).rstrip(']')
            result["title"] = title
            continue
        
        # 요약 파싱
        if line.startswith('요약:'):
            result["summary"] = line.replace('요약:', '').strip()
            continue
        
        # 해시태그 파싱
        if '해시태그:' in line or line.startswith('#'):
            hashtag_text = line.replace('추천 해시태그:', '').replace('해시태그:', '').strip()
            tags = re.findall(r'#\S+', hashtag_text)
            result["hashtags"].extend(tags)
            continue
        
        # FAQ 파싱
        if line.startswith('Q.') or line.startswith('Q:'):
            in_faq = True
            if current_question:
                result["faq"].append(current_question)
            current_question = {"question": line, "answer": ""}
            continue
        
        if in_faq and (line.startswith('A.') or line.startswith('A:')):
            if current_question:
                current_question["answer"] = line
            continue
        
        # 소제목 파싱 ({{소제목}} 또는 **소제목** 또는 ## 소제목)
        heading_match = re.match(r'^(?:\{\{|##|\*\*)\s*(.+?)(?:\}\}|\*\*)?$', line)
        if heading_match or (len(line) < 50 and not line.endswith('.') and not line.startswith('-') and not line.startswith('1')):
            # 이전 섹션 저장
            if current_section:
                current_section["content"] = '\n'.join(current_content).strip()
                if current_section["content"]:
                    result["sections"].append(current_section)
            
            # 새 섹션 시작
            heading = heading_match.group(1) if heading_match else line
            heading = re.sub(r'^[\d]+[\.\)]\s*', '', heading)  # 숫자 제거
            current_section = {"heading": heading, "content": ""}
            current_content = []
            in_faq = False
            continue
        
        # 일반 본문
        if current_section:
            current_content.append(line)
        elif not result["title"]:
            # 첫 번째 줄이 제목
            result["title"] = line
    
    # 마지막 섹션 저장
    if current_section:
        current_section["content"] = '\n'.join(current_content).strip()
        if current_section["content"]:
            result["sections"].append(current_section)
    
    # 마지막 FAQ 저장
    if current_question:
        result["faq"].append(current_question)
    
    return result


def generate_image_prompt(heading, content, main_topic=""):
    """
    섹션 내용 기반 이미지 프롬프트 생성
    
    Args:
        heading: 소제목
        content: 본문 내용
        main_topic: 전체 글의 주제
    
    Returns:
        str: 이미지 생성 프롬프트
    """
    # 핵심 키워드 추출
    keywords = []
    
    # 소제목에서 키워드
    heading_words = re.findall(r'[\w가-힣]+', heading)
    keywords.extend(heading_words[:3])
    
    # 본문에서 키워드 (첫 50자)
    content_preview = content[:100] if content else ""
    content_words = re.findall(r'[\w가-힣]{2,}', content_preview)
    keywords.extend(content_words[:3])
    
    # 중복 제거
    keywords = list(dict.fromkeys(keywords))[:5]
    
    # 프롬프트 생성
    keyword_str = ', '.join(keywords) if keywords else main_topic
    
    prompt = f"Professional blog illustration about {keyword_str}. Modern, clean design, minimalist style, soft colors, suitable for a blog post. No text in the image."
    
    return prompt


def create_word_document(parsed_content, output_path, generate_images=True, image_provider="dalle"):
    """
    파싱된 블로그 콘텐츠로 워드 문서 생성
    
    Args:
        parsed_content: parse_blog_content() 결과
        output_path: 저장 경로
        generate_images: 이미지 생성 여부
        image_provider: 이미지 생성 서비스 (dalle, pollinations)
    
    Returns:
        str: 저장된 파일 경로
    """
    doc = Document()
    
    # 문서 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(11)
    
    # 제목
    if parsed_content["title"]:
        title = doc.add_heading(parsed_content["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 요약
    if parsed_content["summary"]:
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"📌 {parsed_content['summary']}")
        summary_run.bold = True
        summary_run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph()  # 빈 줄
    
    # 섹션별 처리
    image_paths = []
    
    for i, section in enumerate(parsed_content["sections"]):
        heading = section.get("heading", f"섹션 {i+1}")
        content = section.get("content", "")
        
        # 소제목
        doc.add_heading(heading, level=1)
        
        # 이미지 생성 및 삽입
        if generate_images and content:
            print(f"\n[이미지 {i+1}/{len(parsed_content['sections'])}] {heading[:30]}...")
            
            prompt = generate_image_prompt(heading, content, parsed_content["title"])
            
            image_filename = f"section_{i+1:02d}.png"
            image_path = os.path.join(IMAGES_DIR, image_filename)
            
            # 이미지 생성
            result = generate_image(prompt, output_path=image_path, style="blog", provider=image_provider)
            
            if result and os.path.exists(result):
                try:
                    doc.add_picture(result, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    image_paths.append(result)
                except Exception as e:
                    print(f"  ⚠️ 이미지 삽입 실패: {e}")
            
            # API 속도 제한 대응
            time.sleep(2)
        
        # 본문
        if content:
            for para_text in content.split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        
        doc.add_paragraph()  # 섹션 구분
    
    # FAQ
    if parsed_content["faq"]:
        doc.add_heading("자주 묻는 질문", level=1)
        for faq in parsed_content["faq"]:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(faq["question"])
            q_run.bold = True
            
            if faq["answer"]:
                doc.add_paragraph(faq["answer"])
        doc.add_paragraph()
    
    # 해시태그
    if parsed_content["hashtags"]:
        hashtag_para = doc.add_paragraph()
        hashtag_run = hashtag_para.add_run(' '.join(parsed_content["hashtags"]))
        hashtag_run.font.color.rgb = RGBColor(0, 102, 204)
    
    # 저장
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f"\n✓ 워드 문서 저장: {output_path}")
    print(f"  생성된 이미지: {len(image_paths)}개")
    
    return output_path


def process_excel_to_word(excel_path=None, generate_images=True, image_provider="dalle"):
    """
    엑셀 파일의 블로그 본문을 워드 문서로 변환
    
    Args:
        excel_path: 엑셀 파일 경로 (None이면 자동 탐색)
        generate_images: 이미지 생성 여부
        image_provider: 이미지 생성 서비스
    
    Returns:
        list: 생성된 워드 문서 경로 리스트
    """
    if excel_path is None:
        excel_path = find_blog_file()
    
    print(f"엑셀 파일: {excel_path}")
    
    wb = load_workbook(excel_path)
    ws = wb.active
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(IMAGES_DIR, exist_ok=True)
    
    word_files = []
    max_row = ws.max_row
    
    for row in range(2, max_row + 1):
        title = ws[f'A{row}'].value
        content = ws[f'B{row}'].value
        
        if not title or not content:
            continue
        
        print(f"\n{'='*50}")
        print(f"[{row-1}] {title[:40]}...")
        print('='*50)
        
        # 본문 파싱
        parsed = parse_blog_content(content)
        if not parsed["title"]:
            parsed["title"] = title
        
        output_path = os.path.join(OUTPUT_DIR, f"post_{row-1:03d}.docx")
        
        try:
            result = create_word_document(
                parsed, 
                output_path, 
                generate_images=generate_images,
                image_provider=image_provider
            )
            word_files.append(result)
        except Exception as e:
            print(f"  ✗ 워드 생성 오류: {e}")
    
    wb.close()
    
    print(f"\n{'='*50}")
    print(f"완료! 총 {len(word_files)}개 워드 문서 생성")
    print(f"저장 위치: {OUTPUT_DIR}")
    
    return word_files


if __name__ == "__main__":
    import sys
    
    print("=== 워드 문서 생성 (이미지 포함) ===\n")
    
    try:
        excel_file = find_blog_file()
        print(f"엑셀 파일 발견: {excel_file}\n")
        
        result = process_excel_to_word(
            excel_path=excel_file,
            generate_images=True,
            image_provider="dalle"
        )
        
        print(f"\n생성된 워드 문서: {len(result)}개")
        for f in result:
            print(f"  - {f}")
            
    except FileNotFoundError:
        print("blog*.xlsx 파일을 찾을 수 없습니다.")
        print("먼저 '엑셀 파일 생성'과 '블로그 본문 생성'을 실행하세요.")
        sys.exit(1)
    except Exception as e:
        print(f"오류 발생: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
